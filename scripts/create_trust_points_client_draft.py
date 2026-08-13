from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "Tshelo_Trust_Points_Client_Review_Draft.docx"

PURPLE = "7439E0"
PURPLE_DARK = "4E229A"
PURPLE_LIGHT = "EEE7FC"
PURPLE_PALE = "F8F5FE"
INK = "17121F"
MUTED = "66616F"
GRAY = "E3E0E7"
LIGHT_GRAY = "F4F3F6"
GREEN = "087B5B"
AMBER = "916300"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=GRAY, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_font(run, size=None, bold=None, color=INK, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def format_cell_text(cell, size=9.2, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            set_font(run, size=size, color=color, bold=bold)


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))


def add_body(doc, text, bold_lead=None, after=6, color=INK, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, size=11, bold=True, color=color)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest, size=11, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_font(run, size=11, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_font(r, size=10.8, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_font(r, size=10.8, color=INK)
    return p


def add_callout(doc, label, text, fill=PURPLE_PALE, accent=PURPLE_DARK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size=8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.10
    lead = p.add_run(label.upper() + "  ")
    set_font(lead, size=9.5, bold=True, color=accent)
    body = p.add_run(text)
    set_font(body, size=10.5, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, font_size=9.0, header_fill=PURPLE_LIGHT, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        hdr.cells[i].text = header
        set_cell_shading(hdr.cells[i], header_fill)
        format_cell_text(hdr.cells[i], size=9.1, color=PURPLE_DARK, bold=True,
                         align=(alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT))
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = str(value)
            if row_index % 2 == 1:
                set_cell_shading(cells[i], "FBFAFC")
            format_cell_text(cells[i], size=font_size, color=INK,
                             align=(alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT))
    set_table_geometry(table, widths)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, 16, 8, PURPLE_DARK),
        "Heading 2": (13, 12, 6, PURPLE_DARK),
        "Heading 3": (12, 8, 4, PURPLE_DARK),
    }
    for style_name, (size, before, after, color) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("TSHELO  |  CLIENT REVIEW DRAFT")
    set_font(hr, size=8.5, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    left = fp.add_run("Trust Points & User Rewards  •  12 August 2026")
    set_font(left, size=8.5, color=MUTED)
    fp.add_run("\t")
    add_page_number(fp)

    # Page 1 — decision brief
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(28)
    kicker.paragraph_format.space_after = Pt(7)
    kr = kicker.add_run("CLIENT DECISION BRIEF")
    set_font(kr, size=10, bold=True, color=PURPLE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(7)
    tr = title.add_run("Tshelo Trust Points\nand User Rewards")
    set_font(tr, size=28, bold=True, color=INK)
    set_paragraph_keep(title)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    sr = subtitle.add_run("Draft for client verification, approval, or amendment")
    set_font(sr, size=14, color=MUTED)

    meta = add_table(
        doc,
        ["Document status", "Prepared for", "Review date"],
        [["DRAFT — not final policy", "Tshelo client team", "12 August 2026"]],
        [2600, 3800, 2960],
        font_size=9.5,
        header_fill=LIGHT_GRAY,
    )

    add_callout(
        doc,
        "Decision requested",
        "Please approve the point schedule, trust-level thresholds, user benefits, and deduction rules. Any item marked PROPOSED is a recommendation and can be changed before launch.",
    )

    add_heading(doc, "What this document covers", 1)
    add_body(doc, "Tshelo uses trust points to recognise reliable participation, transparent organising, and responsible contribution behaviour. This draft explains what is already implemented and proposes what users should receive at each trust level.")
    add_bullet(doc, "How users earn the full 100 trust points.")
    add_bullet(doc, "How New, Basic, Trusted, and Verified levels should be assigned.")
    add_bullet(doc, "Which badges, limits, service benefits, and optional commercial rewards each level should receive.")
    add_bullet(doc, "How reports, flags, appeals, Rich Auntie recognition, and paid tokens should be handled.")

    add_heading(doc, "Recommended direction", 1)
    add_body(doc, "Launch with non-cash benefits first: visible trust badges, progressively higher fund limits, appropriate profile signals, and support priority. Keep free exports, token discounts, and other cost-bearing benefits optional until their commercial impact is approved.")
    add_callout(doc, "Core principle", "Trust points are reputation, not money. They cannot be bought, transferred, exchanged for tokens, redeemed, or converted to cash.", fill="F3FBF8", accent=GREEN)

    # Page 2 — implemented mechanics and schedule
    add_page_break(doc)
    add_heading(doc, "1. Current trust model", 1)
    add_body(doc, "The following mechanics are already represented in the current product and backend. The client is being asked to verify that these rules reflect the intended policy.")
    add_number(doc, "The server evaluates completed actions and awards each achievement once.")
    add_number(doc, "Each achievement contributes a fixed number of trust points.")
    add_number(doc, "The achievement path totals 100 points before any deductions.")
    add_number(doc, "The app shows the score out of 100, the trust level, earned achievements, and progress toward incomplete achievements.")
    add_number(doc, "A user receives an in-app notification when an achievement is unlocked or their trust level changes.")

    add_heading(doc, "Achievement and point schedule", 2)
    achievement_rows = [
        ("Profile Ready", "Complete the Tshelo profile", "5"),
        ("Payment Identity Verified", "Verify the registered mobile-money identity", "5"),
        ("First Contribution", "Make a first confirmed contribution", "5"),
        ("Consistent Contributor", "Contribute to 3 different funds", "10"),
        ("Community Pillar", "Contribute to 10 different funds", "15"),
        ("Receipt Starter", "Add valid receipts to 3 expenses", "5"),
        ("Transparent Organiser", "Maintain at least 80% receipt coverage across 5+ expenses in one fund", "15"),
        ("First Fund Completed", "Successfully close a first fund", "10"),
        ("Reliable Organiser", "Close 3 funds without unresolved disputes", "15"),
        ("Goal Getter", "Lead a fund that reaches 100% of its target", "10"),
        ("Event Ready", "Complete an event date, time, venue, and initial guest list", "5"),
        ("TOTAL", "Complete achievement path", "100"),
    ]
    table = add_table(
        doc,
        ["Achievement", "Requirement", "Points"],
        achievement_rows,
        [2550, 5790, 1020],
        font_size=8.8,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )
    for cell in table.rows[-1].cells:
        set_cell_shading(cell, PURPLE_LIGHT)
        format_cell_text(cell, size=9, color=PURPLE_DARK, bold=True,
                         align=WD_ALIGN_PARAGRAPH.CENTER if cell == table.rows[-1].cells[-1] else WD_ALIGN_PARAGRAPH.LEFT)
    add_callout(doc, "Client check", "Approve the achievement names, requirements, and point values above, or mark the changes required in the approval section.", fill="FFF9EA", accent=AMBER)

    # Page 3 — levels
    add_page_break(doc)
    add_heading(doc, "2. Trust levels and eligibility", 1)
    add_body(doc, "For clarity and fairness, this draft recommends that trust levels follow the score bands below. Verified status also requires mobile-money identity verification.")
    level_rows = [
        ("New", "0–29", "No additional requirement", "Early-stage or limited verified activity"),
        ("Basic", "30–59", "None", "Established activity and initial reliability"),
        ("Trusted", "60–79", "None", "Strong history of contribution or organising"),
        ("Verified", "80–100", "Mobile-money identity verified", "Highest trust level with verified payment identity"),
    ]
    add_table(
        doc,
        ["Level", "Score", "Additional requirement", "Meaning"],
        level_rows,
        [1350, 1150, 2700, 4160],
        font_size=9.0,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    add_heading(doc, "Existing exception requiring a client decision", 2)
    add_body(doc, "The current backend also allows a user to become Basic after completing 1 fund, or Trusted after completing 3 funds, even when the score is below the normal band.")
    add_callout(doc, "Recommendation", "Remove the completed-fund shortcut. Fund completion already earns trust points through First Fund Completed and Reliable Organiser. Keeping both rules can double-count the same behaviour and make levels harder to explain.", fill="FFF4F4", accent=RED)

    add_heading(doc, "What the user sees", 2)
    add_bullet(doc, "A trust score from 0 to 100 and a progress bar.")
    add_bullet(doc, "A label: New, Basic, Trusted, or Verified.")
    add_bullet(doc, "Earned achievement cards and progress toward the remaining achievements.")
    add_bullet(doc, "A short explanation that trust points have no cash value and are separate from paid tokens.")
    add_bullet(doc, "Rich Auntie recognition as a separate organiser-awarded honour.")

    add_heading(doc, "Level changes", 2)
    add_body(doc, "Levels should update automatically when achievements are awarded or an approved deduction is applied. The user should be notified of the new level and should be able to understand which completed achievements and deductions produced the score.")

    # Page 4 — rewards proposal
    add_page_break(doc)
    add_heading(doc, "3. Proposed user rewards by trust level", 1)
    add_body(doc, "The benefits below are proposed for client approval. The launch recommendation favours benefits that strengthen trust and product access without creating a cash or token liability.")

    reward_rows = [
        ("New", "Progress tracker and achievement notifications; standard app access; default fund limits.", "None at launch."),
        ("Basic", "Visible Basic badge; modestly higher country-specific fund-goal and active-fund limits; access to standard organiser templates.", "Occasional educational tips or guided organiser checklist."),
        ("Trusted", "Visible Trusted badge; higher fund-goal and active-fund limits; priority support queue; eligibility for selected pilot features.", "One complimentary premium PDF export per month, subject to cost approval."),
        ("Verified", "Verified shield; highest approved fund limits; verified identity shown where members select organisers/admins; highest support priority; early feature access.", "Two complimentary premium PDF exports per month or a small token discount, subject to commercial approval."),
    ]
    add_table(
        doc,
        ["Level", "Recommended launch benefits", "Optional later benefit"],
        reward_rows,
        [1250, 5280, 2830],
        font_size=8.8,
    )

    add_heading(doc, "Benefits that should not be granted", 2)
    add_bullet(doc, "Cash, money-equivalent value, or withdrawal rights.")
    add_bullet(doc, "Automatic paid-token credits for achievements.")
    add_bullet(doc, "Permission to bypass identity, payment, fraud, dispute, or admin controls.")
    add_bullet(doc, "Guaranteed prominence that hides lower-trust users or makes trust appear purchased.")
    add_bullet(doc, "Permanent benefits that cannot be removed after an upheld report or policy breach.")

    add_heading(doc, "Recommended launch choice", 2)
    add_callout(doc, "Phase 1", "Approve badges, progressive country-specific fund limits, support priority, and verified identity display. Defer free exports and token discounts until pricing, gateway costs, taxes, and expected usage are approved.")

    add_heading(doc, "Fund-limit values", 2)
    add_body(doc, "Exact maximum fund goals and active-fund counts should be set per country and per trust level. These figures should be approved separately because currency, risk, and operational capacity differ by market.")

    # Page 5 — safeguards and client approval form
    add_page_break(doc)
    add_heading(doc, "4. Deductions, safeguards, and related recognition", 1)
    deduction_rows = [
        ("Reported fund", "Current rule: −10 points per reported fund", "Apply only after review confirms the report; do not deduct for an unverified complaint."),
        ("Flagged account", "Current rule: −20 points", "Require a recorded reason, reviewer, date, and appeal/review path."),
        ("Score floor", "Minimum 0", "Never show a negative trust score."),
        ("Score ceiling", "Maximum 100", "Additional achievements should not inflate the score beyond 100."),
    ]
    add_table(
        doc,
        ["Condition", "Current calculation", "Recommended policy safeguard"],
        deduction_rows,
        [2100, 2400, 4860],
        font_size=9.0,
    )

    add_heading(doc, "Required controls", 2)
    add_bullet(doc, "Awards are server-owned and granted once; the app cannot credit its own points.")
    add_bullet(doc, "Every score-changing action should be auditable.")
    add_bullet(doc, "Users should be able to see deductions and request review.")
    add_bullet(doc, "Deleting or reversing qualifying activity should trigger a defined re-evaluation policy.")
    add_bullet(doc, "Trust status must never replace payment verification, admin permissions, or fraud checks.")

    add_heading(doc, "Separate concepts", 2)
    add_body(doc, "Paid tokens:", bold_lead="Paid tokens:")
    add_body(doc, "Purchased app credit used for clearly priced premium actions. Tokens are not trust points.", after=4)
    add_body(doc, "Rich Auntie recognition:", bold_lead="Rich Auntie recognition:")
    add_body(doc, "A special appreciation award given by organisers for meaningful sponsorship or support. It remains visible and celebratory but does not automatically add trust points.", after=4)

    add_heading(doc, "5. Client approval checklist", 1)
    approval_rows = [
        ("Achievement schedule totals 100 points", "☐ Approve  ☐ Change", ""),
        ("Level bands: 0–29 / 30–59 / 60–79 / 80–100", "☐ Approve  ☐ Change", ""),
        ("Remove completed-fund shortcuts for Basic/Trusted", "☐ Approve  ☐ Keep", ""),
        ("Phase 1 rewards: badges, limits, support priority, verified identity", "☐ Approve  ☐ Change", ""),
        ("Defer free exports and token discounts", "☐ Approve  ☐ Include", ""),
        ("Apply deductions only after review, with appeal visibility", "☐ Approve  ☐ Change", ""),
        ("Keep Rich Auntie separate from trust points", "☐ Approve  ☐ Change", ""),
    ]
    add_table(
        doc,
        ["Decision", "Client selection", "Required changes / notes"],
        approval_rows,
        [4500, 2200, 2660],
        font_size=8.7,
        header_fill=LIGHT_GRAY,
    )

    add_heading(doc, "Approval", 2)
    add_body(doc, "Client representative:  ______________________________________________")
    add_body(doc, "Decision:  ☐ Approved as drafted   ☐ Approved with changes   ☐ Further review required")
    add_body(doc, "Signature:  ____________________________________     Date:  __________________")
    add_body(doc, "Additional comments:")
    for _ in range(3):
        p = doc.add_paragraph("________________________________________________________________________________")
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            set_font(run, size=10, color=GRAY)

    properties = doc.core_properties
    properties.title = "Tshelo Trust Points and User Rewards — Client Review Draft"
    properties.subject = "Client verification and approval of Tshelo trust points, levels, rewards, and safeguards"
    properties.author = "Tshelo Product Team"
    properties.keywords = "Tshelo, trust points, rewards, client approval, draft"
    properties.comments = "Draft for client review; not final policy."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
